import os
import uuid
import shutil
import docx
from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename
from dotenv import load_dotenv

# Load env variables from .env
load_dotenv()

# Langchain and vector store imports
from langchain_core.documents import Document
from langchain_community.document_loaders.pdf import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_groq import ChatGroq
from sentence_transformers import SentenceTransformer
import chromadb

app = Flask(__name__)

# Config
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_BASE_FOLDER = os.path.join(BASE_DIR, 'data', 'uploads')
PERSIST_DIRECTORY = os.path.join(BASE_DIR, 'data', 'vector_store')
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}

app.config['UPLOAD_BASE_FOLDER'] = UPLOAD_BASE_FOLDER
app.config['TEMPLATES_AUTO_RELOAD'] = True
os.makedirs(UPLOAD_BASE_FOLDER, exist_ok=True)

# ----------------------------------------------------
# Singleton Model & Vector Store Managers
# ----------------------------------------------------
class EmbeddingManager:
    def __init__(self, model_name="all-MiniLM-L6-v2"):
        self.model_name = model_name
        self.model = None   # do not load at startup

    def _load_model(self):
        if self.model is None:
            print(f"[*] Loading embedding model: {self.model_name}...")
            self.model = SentenceTransformer(self.model_name)
            print("[*] Embedding model loaded successfully.")

    def generate_embeddings(self, texts):
        self._load_model()
        return self.model.encode(texts, show_progress_bar=False)


class VectorStoreManager:
    def __init__(self, persist_dir=PERSIST_DIRECTORY):
        self.persist_directory = persist_dir
        os.makedirs(self.persist_directory, exist_ok=True)
        self.client = chromadb.PersistentClient(path=self.persist_directory)
        print(f"[*] Vector store client initialized at: {self.persist_directory}")

    def _get_collection_name(self, session_id):
        # Format session ID into a safe collection name (alphanumeric, underscores, 3-63 chars)
        safe_id = str(session_id).replace('-', '_')
        return f"col_{safe_id}"[:63]

    def get_collection(self, session_id):
        col_name = self._get_collection_name(session_id)
        return self.client.get_or_create_collection(
            name=col_name,
            metadata={"description": f"Vector store collection for user session {session_id}"}
        )

    def add_documents(self, session_id, documents, embeddings):
        if not documents:
            return
        
        collection = self.get_collection(session_id)
        ids = []
        metadatas = []
        documents_content = []
        embeddings_list = []

        for i, (doc, embedding) in enumerate(zip(documents, embeddings)):
            doc_id = f"doc_{uuid.uuid4()}"
            ids.append(doc_id)

            metadata = dict(doc.metadata)
            metadata["doc_index"] = i
            metadata["content_length"] = len(doc.page_content)
            if "source" in metadata:
                metadata["source_filename"] = os.path.basename(metadata["source"])
            else:
                metadata["source_filename"] = "unknown"

            metadatas.append(metadata)
            documents_content.append(doc.page_content)
            embeddings_list.append(embedding.tolist())

        # Insert in a single batch
        collection.add(
            ids=ids,
            metadatas=metadatas,
            documents=documents_content,
            embeddings=embeddings_list
        )
        print(f"[*] Added {len(documents)} document chunks to collection {collection.name}. Count: {collection.count()}")

    def get_indexed_files(self, session_id):
        collection = self.get_collection(session_id)
        all_data = collection.get(include=["metadatas"])
        if not all_data or not all_data.get("metadatas"):
            return []
        
        filenames = set()
        for meta in all_data["metadatas"]:
            filename = meta.get("source_filename")
            if filename:
                filenames.add(filename)
        return sorted(list(filenames))

    def reset_collection(self, session_id):
        col_name = self._get_collection_name(session_id)
        try:
            self.client.delete_collection(col_name)
            print(f"[*] Deleted collection: {col_name}")
        except Exception as e:
            print(f"[!] Warning deleting collection {col_name}: {e}")
        
        # Trigger recreation
        self.get_collection(session_id)


# Initialize managers
embedding_manager = EmbeddingManager()
vector_store = VectorStoreManager()

# ----------------------------------------------------
# Document Ingestion Helpers
# ----------------------------------------------------
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def load_docx_file(file_path):
    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
                    
    full_text = "\n\n".join(paragraphs)
    return [Document(page_content=full_text, metadata={"source": file_path, "page": 1})]

def load_txt_file(file_path):
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [Document(page_content=text, metadata={"source": file_path, "page": 1})]

def ingest_file(session_id, file_path, file_extension):
    documents = []
    
    if file_extension == 'pdf':
        loader = PyPDFLoader(file_path)
        documents = loader.load()
    elif file_extension == 'docx':
        documents = load_docx_file(file_path)
    elif file_extension == 'txt':
        documents = load_txt_file(file_path)
        
    if not documents:
        return 0

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    chunks = text_splitter.split_documents(documents)

    if not chunks:
        return 0

    texts = [chunk.page_content for chunk in chunks]
    embeddings = embedding_manager.generate_embeddings(texts)
    vector_store.add_documents(session_id, chunks, embeddings)
    
    return len(chunks)

# ----------------------------------------------------
# Flask Routes
# ----------------------------------------------------
@app.route('/')
def home():
    return render_template('index.html')

@app.route('/api/files', methods=['GET'])
def get_files():
    session_id = request.headers.get('X-Session-ID')
    if not session_id:
        return jsonify({"success": False, "error": "Missing X-Session-ID header"}), 400
        
    files = vector_store.get_indexed_files(session_id)
    return jsonify({"success": True, "files": files})

@app.route('/api/upload', methods=['POST'])
def upload_file():
    session_id = (
        request.headers.get('X-Session-ID')
        or (request.get_json(silent=True) or {}).get('session_id')
        or request.form.get('session_id')
    )
    if not session_id:
        return jsonify({"success": False, "error": "Missing X-Session-ID header"}), 400

    if 'file' not in request.files:
        return jsonify({"success": False, "error": "No file part in the request"}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({"success": False, "error": "No file selected"}), 400

    if not allowed_file(file.filename):
        return jsonify({"success": False, "error": f"File type not supported. Allowed formats: {', '.join(ALLOWED_EXTENSIONS)}"}), 400

    try:
        # Create user-specific upload directory
        user_upload_dir = os.path.join(app.config['UPLOAD_BASE_FOLDER'], secure_filename(session_id))
        os.makedirs(user_upload_dir, exist_ok=True)
        
        filename = secure_filename(file.filename)
        file_path = os.path.join(user_upload_dir, filename)
        file.save(file_path)

        file_ext = filename.rsplit('.', 1)[1].lower()
        chunks_added = ingest_file(session_id, file_path, file_ext)

        return jsonify({
            "success": True, 
            "filename": filename, 
            "chunks": chunks_added,
            "message": f"Successfully uploaded and indexed '{filename}' ({chunks_added} chunks)."
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": f"Error indexing file: {str(e)}"}), 500

@app.route('/api/query', methods=['POST'])
def query_rag():
    session_id = request.headers.get('X-Session-ID')
    if not session_id:
        return jsonify({"success": False, "error": "Missing X-Session-ID header"}), 400

    data = request.json or {}
    question = data.get('question')
    show_references = data.get('show_references', False)

    if not question or not question.strip():
        return jsonify({"success": False, "error": "Question cannot be empty"}), 400

    collection = vector_store.get_collection(session_id)
    if collection.count() == 0:
        return jsonify({
            "success": True,
            "answer": "Please upload and index some documents first before asking questions.",
            "references": []
        })

    try:
        # 1. Generate Query Embeddings
        query_embeddings = embedding_manager.generate_embeddings([question])[0]

        # 2. Query Chroma Vector store (retrieve top 5 results)
        results = collection.query(
            query_embeddings=[query_embeddings.tolist()],
            n_results=5
        )

        retrieved_docs = []
        if results and results.get("documents") and results["documents"][0]:
            ids = results["ids"][0]
            metadatas = results["metadatas"][0]
            documents = results["documents"][0]
            distances = results["distances"][0]

            for i, (doc_id, metadata, doc_text, distance) in enumerate(zip(ids, metadatas, documents, distances)):
                similarity_score = float(1.0 - distance)
                source = metadata.get("source_filename", "unknown")
                page = metadata.get("page", 1)
                
                retrieved_docs.append({
                    "id": doc_id,
                    "text": doc_text,
                    "source": source,
                    "page": page,
                    "similarity": round(similarity_score, 4)
                })

        # 3. Form Context & Prompt
        context = "\n\n".join([f"Source: {doc['source']} (Page {doc['page']})\nContent: {doc['text']}" for doc in retrieved_docs])

        prompt = f"""Use the following retrieved context segments from documents to answer the question.
If the context doesn't contain enough information to answer, state that you cannot find the answer in the uploaded files. 
Be concise, factually accurate, and prioritize information found in the context.

Context:
{context}

Question:
{question}

Answer:"""

        # 4. Connect to Groq
        api_key = os.environ.get("GROQ_API_KEY")
        model_name = os.environ.get("GROQ_MODEL", "llama-3.1-8b-instant")
        
        if not api_key:
            return jsonify({"success": False, "error": "GROQ_API_KEY is not configured in"}), 500

        llm = ChatGroq(
            groq_api_key=api_key,
            model=model_name,
            temperature=0.1,
            max_tokens=1024
        )

        # Call LLM
        response = llm.invoke(prompt)
        answer = response.content

        references = []
        if show_references:
            references = retrieved_docs

        return jsonify({
            "success": True,
            "answer": answer,
            "references": references
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": f"Error running query: {str(e)}"}), 500

@app.route('/api/clear', methods=['POST'])
def clear_db():
    session_id = request.headers.get('X-Session-ID')
    if not session_id:
        return jsonify({"success": False, "error": "Missing X-Session-ID header"}), 400

    try:
        # Clear vector database for this user
        vector_store.reset_collection(session_id)
        
        # Delete upload files for this user
        user_upload_dir = os.path.join(app.config['UPLOAD_BASE_FOLDER'], secure_filename(session_id))
        if os.path.exists(user_upload_dir):
            shutil.rmtree(user_upload_dir)

        return jsonify({"success": True, "message": "Database and uploaded documents cleared successfully."})
    except Exception as e:
        return jsonify({"success": False, "error": f"Error resetting database: {str(e)}"}), 500


if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(debug=False, host='0.0.0.0', port=port)
